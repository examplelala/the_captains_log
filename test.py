from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.shared import OxmlElement, qn
from docx.oxml import parse_xml

def set_cell_background_color(cell, color_hex):
    """设置表格单元格背景色"""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), color_hex)
    tcPr.append(shd)

def set_cell_borders(cell, top=None, bottom=None, left=None, right=None):
    """设置单元格边框"""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = tcPr.find(qn('w:tcBorders'))
    if tcBorders is None:
        tcBorders = OxmlElement('w:tcBorders')
        tcPr.append(tcBorders)
    
    borders = {'top': top, 'bottom': bottom, 'left': left, 'right': right}
    for border_name, border_config in borders.items():
        if border_config:
            border_elem = tcBorders.find(qn(f'w:{border_name}'))
            if border_elem is None:
                border_elem = OxmlElement(f'w:{border_name}')
                tcBorders.append(border_elem)
            border_elem.set(qn('w:val'), border_config.get('style', 'single'))
            border_elem.set(qn('w:sz'), str(border_config.get('width', 4)))
            border_elem.set(qn('w:color'), border_config.get('color', '000000'))

def create_professional_resume():
    # 创建文档
    doc = Document()
    
    # 设置页面边距 - 更紧凑的布局
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(0.4)
        section.bottom_margin = Inches(0.4)
        section.left_margin = Inches(0.3)
        section.right_margin = Inches(0.3)
    
    # ==================== 创建主布局表格 ====================
    # 左右两栏布局：左栏放个人信息和技能，右栏放工作经历等
    main_table = doc.add_table(rows=1, cols=2)
    main_table.autofit = False
    main_table.columns[0].width = Inches(2.6)  # 左栏宽度
    main_table.columns[1].width = Inches(5.4)  # 右栏宽度
    
    left_cell = main_table.cell(0, 0)
    right_cell = main_table.cell(0, 1)
    
    # 左栏设置深色背景
    set_cell_background_color(left_cell, '2C3E50')  # 深蓝灰色
    # 右栏设置浅色背景
    set_cell_background_color(right_cell, 'FFFFFF')  # 白色
    
    # 设置单元格垂直对齐
    left_cell.vertical_alignment = WD_ALIGN_VERTICAL.TOP
    right_cell.vertical_alignment = WD_ALIGN_VERTICAL.TOP
    
    # 移除边框
    for cell in [left_cell, right_cell]:
        set_cell_borders(cell, 
                        top={'style': 'nil'},
                        bottom={'style': 'nil'},
                        left={'style': 'nil'},
                        right={'style': 'nil'})
    
    # ==================== 左栏内容 ====================
    # 清空左栏默认段落
    left_cell.paragraphs[0].clear()
    
    # 头像区域
    avatar_para = left_cell.add_paragraph()
    avatar_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    avatar_para.space_after = Pt(20)
    avatar_run = avatar_para.add_run('头像位置')
    avatar_run.font.size = Pt(10)
    avatar_run.font.color.rgb = RGBColor(189, 195, 199)  # 浅灰色
    
    # 姓名
    name_para = left_cell.add_paragraph()
    name_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    name_para.space_after = Pt(8)
    name_run = name_para.add_run('何晓珊')
    name_run.font.size = Pt(24)
    name_run.font.bold = True
    name_run.font.color.rgb = RGBColor(255, 255, 255)  # 白色
    
    # 职位
    job_para = left_cell.add_paragraph()
    job_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    job_para.space_after = Pt(25)
    job_run = job_para.add_run('小学数学老师')
    job_run.font.size = Pt(12)
    job_run.font.color.rgb = RGBColor(52, 152, 219)  # 亮蓝色
    
    # 联系信息标题
    contact_title = left_cell.add_paragraph()
    contact_title.space_before = Pt(10)
    contact_title.space_after = Pt(15)
    contact_title_run = contact_title.add_run('联系方式')
    contact_title_run.font.size = Pt(14)
    contact_title_run.font.bold = True
    contact_title_run.font.color.rgb = RGBColor(255, 255, 255)
    
    # 添加联系信息下划线
    underline_para = left_cell.add_paragraph()
    underline_para.space_after = Pt(15)
    underline_run = underline_para.add_run('_' * 25)
    underline_run.font.color.rgb = RGBColor(52, 152, 219)
    
    # 联系信息详情
    contacts = [
        ('📍', 'China'),
        ('📱', '+86151-1197-5608'),
        ('📧', '1294028360@qq.com')
    ]
    
    for icon, info in contacts:
        contact_para = left_cell.add_paragraph()
        contact_para.space_after = Pt(8)
        # 图标
        icon_run = contact_para.add_run(f'{icon} ')
        icon_run.font.size = Pt(12)
        # 信息
        info_run = contact_para.add_run(info)
        info_run.font.size = Pt(10)
        info_run.font.color.rgb = RGBColor(189, 195, 199)
    
    # 技能与荣誉部分
    skills_title = left_cell.add_paragraph()
    skills_title.space_before = Pt(25)
    skills_title.space_after = Pt(15)
    skills_title_run = skills_title.add_run('技能与荣誉')
    skills_title_run.font.size = Pt(14)
    skills_title_run.font.bold = True
    skills_title_run.font.color.rgb = RGBColor(255, 255, 255)
    
    # 技能下划线
    skills_underline = left_cell.add_paragraph()
    skills_underline.space_after = Pt(15)
    skills_underline_run = skills_underline.add_run('_' * 25)
    skills_underline_run.font.color.rgb = RGBColor(52, 152, 219)
    
    # 专业证书
    cert_title = left_cell.add_paragraph()
    cert_title.space_after = Pt(5)
    cert_title_run = cert_title.add_run('专业证书')
    cert_title_run.font.size = Pt(11)
    cert_title_run.font.bold = True
    cert_title_run.font.color.rgb = RGBColor(52, 152, 219)
    
    cert_content = left_cell.add_paragraph()
    cert_content.space_after = Pt(15)
    cert_content_run = cert_content.add_run('• 小学数学教师资格证\n• 普通话二级甲等证书')
    cert_content_run.font.size = Pt(9)
    cert_content_run.font.color.rgb = RGBColor(189, 195, 199)
    
    # 语言能力
    lang_title = left_cell.add_paragraph()
    lang_title.space_after = Pt(5)
    lang_title_run = lang_title.add_run('语言能力')
    lang_title_run.font.size = Pt(11)
    lang_title_run.font.bold = True
    lang_title_run.font.color.rgb = RGBColor(52, 152, 219)
    
    lang_content = left_cell.add_paragraph()
    lang_content.space_after = Pt(15)
    lang_content_run = lang_content.add_run('• CET6\n• CET4')
    lang_content_run.font.size = Pt(9)
    lang_content_run.font.color.rgb = RGBColor(189, 195, 199)
    
    # 获得奖项
    award_title = left_cell.add_paragraph()
    award_title.space_after = Pt(5)
    award_title_run = award_title.add_run('主要奖项')
    award_title_run.font.size = Pt(11)
    award_title_run.font.bold = True
    award_title_run.font.color.rgb = RGBColor(52, 152, 219)
    
    award_content = left_cell.add_paragraph()
    award_content_run = award_content.add_run('• 硕士学业二等奖学金（省级）\n• 硕士校长奖学金\n• 教学技能大赛三等奖\n• "挑战杯"多项获奖\n• 创新创业大赛一等奖')
    award_content_run.font.size = Pt(9)
    award_content_run.font.color.rgb = RGBColor(189, 195, 199)
    
    # ==================== 右栏内容 ====================
    # 清空右栏默认段落
    right_cell.paragraphs[0].clear()
    
    # 教育经历
    edu_title = right_cell.add_paragraph()
    edu_title.space_before = Pt(15)
    edu_title.space_after = Pt(5)
    edu_title_run = edu_title.add_run('教育经历')
    edu_title_run.font.size = Pt(16)
    edu_title_run.font.bold = True
    edu_title_run.font.color.rgb = RGBColor(44, 62, 80)
    
    # 教育经历下划线
    edu_line = right_cell.add_paragraph()
    edu_line.space_after = Pt(15)
    edu_line_run = edu_line.add_run('━' * 45)
    edu_line_run.font.color.rgb = RGBColor(52, 152, 219)
    edu_line_run.font.size = Pt(8)
    
    # 学历信息
    edu_school = right_cell.add_paragraph()
    edu_school.space_after = Pt(3)
    edu_school_run = edu_school.add_run('四川师范大学')
    edu_school_run.font.size = Pt(13)
    edu_school_run.font.bold = True
    edu_school_run.font.color.rgb = RGBColor(44, 62, 80)
    
    edu_degree = right_cell.add_paragraph()
    edu_degree.space_after = Pt(3)
    edu_degree_run = edu_degree.add_run('小学教育（本硕一致） | 2023.09-2026.06')
    edu_degree_run.font.size = Pt(11)
    edu_degree_run.font.color.rgb = RGBColor(127, 140, 141)
    
    edu_courses = right_cell.add_paragraph()
    edu_courses.space_after = Pt(25)
    edu_courses.paragraph_format.left_indent = Inches(0.2)
    edu_courses_run = edu_courses.add_run('主修课程：课程与教学论；小学数学教学设计与案例分析；小学德育与班主任工作；小学德育原理学等')
    edu_courses_run.font.size = Pt(10)
    edu_courses_run.font.color.rgb = RGBColor(52, 73, 94)
    
    # 工作经历
    work_title = right_cell.add_paragraph()
    work_title.space_after = Pt(5)
    work_title_run = work_title.add_run('工作经历')
    work_title_run.font.size = Pt(16)
    work_title_run.font.bold = True
    work_title_run.font.color.rgb = RGBColor(44, 62, 80)
    
    # 工作经历下划线
    work_line = right_cell.add_paragraph()
    work_line.space_after = Pt(15)
    work_line_run = work_line.add_run('━' * 45)
    work_line_run.font.color.rgb = RGBColor(52, 152, 219)
    work_line_run.font.size = Pt(8)
    
    # 工作经历1
    work1_company = right_cell.add_paragraph()
    work1_company.space_after = Pt(2)
    work1_company_run = work1_company.add_run('成都市四川师范大学附属实验小学')
    work1_company_run.font.size = Pt(12)
    work1_company_run.font.bold = True
    work1_company_run.font.color.rgb = RGBColor(44, 62, 80)
    
    work1_position = right_cell.add_paragraph()
    work1_position.space_after = Pt(8)
    work1_position_run = work1_position.add_run('实习数学老师 | 2025.03-2025.06')
    work1_position_run.font.size = Pt(10)
    work1_position_run.font.color.rgb = RGBColor(127, 140, 141)
    
    # 工作内容1 - 使用项目符号
    work_items1 = [
        '教学实践：独立设计实施4节新课（含2节综合实践课）；独立撰写70余篇教案',
        '学生辅导：组织4次练习课与5次自习课，运用动态量化表追踪90+学生作业及试卷质量，提升批改效率',
        '教研成长：参与10+次集体备课及评课；随堂听课100+节，观摩公开课10+节，提炼课堂互动技巧',
        '班级管理：协助午餐、早操及运动会秩序管理，有效维持纪律，增强团队沟通与合作意识'
    ]
    
    for item in work_items1:
        work_item_para = right_cell.add_paragraph()
        work_item_para.space_after = Pt(4)
        work_item_para.paragraph_format.left_indent = Inches(0.2)
        work_item_run = work_item_para.add_run(f'• {item}')
        work_item_run.font.size = Pt(9)
        work_item_run.font.color.rgb = RGBColor(52, 73, 94)
    
    # 添加间距
    right_cell.add_paragraph().space_after = Pt(15)
    
    # 工作经历2
    work2_company = right_cell.add_paragraph()
    work2_company.space_after = Pt(2)
    work2_company_run = work2_company.add_run('重庆市南岸区长生桥小学校')
    work2_company_run.font.size = Pt(12)
    work2_company_run.font.bold = True
    work2_company_run.font.color.rgb = RGBColor(44, 62, 80)
    
    work2_position = right_cell.add_paragraph()
    work2_position.space_after = Pt(8)
    work2_position_run = work2_position.add_run('实习数学老师 | 2023.03-2023.06')
    work2_position_run.font.size = Pt(10)
    work2_position_run.font.color.rgb = RGBColor(127, 140, 141)
    
    # 工作内容2
    work_items2 = [
        '教学实践：独立完成数学授课2节，兼任美术、道法等学科教学；独立撰写教案30+篇',
        '学生辅导：日均跟班管理8小时，累计提供60小时+课后服务；负责作业批改与分类指导，协助教师答疑',
        '教研成长：随堂听课100+节，观摩优秀教师公开课10+节，参与集体备课与评课5次',
        '班级管理：协助策划并组织3场主题班会活动，独立管理自习课10+节，持续维护班级秩序与活动纪律'
    ]
    
    for item in work_items2:
        work_item_para = right_cell.add_paragraph()
        work_item_para.space_after = Pt(4)
        work_item_para.paragraph_format.left_indent = Inches(0.2)
        work_item_run = work_item_para.add_run(f'• {item}')
        work_item_run.font.size = Pt(9)
        work_item_run.font.color.rgb = RGBColor(52, 73, 94)
    
    # 添加间距
    right_cell.add_paragraph().space_after = Pt(15)
    
    # 工作经历3
    work3_company = right_cell.add_paragraph()
    work3_company.space_after = Pt(2)
    work3_company_run = work3_company.add_run('成都市四川师范大学附属实验小学')
    work3_company_run.font.size = Pt(12)
    work3_company_run.font.bold = True
    work3_company_run.font.color.rgb = RGBColor(44, 62, 80)
    
    work3_position = right_cell.add_paragraph()
    work3_position.space_after = Pt(8)
    work3_position_run = work3_position.add_run('见习数学老师 | 2024.10-2024.11')
    work3_position_run.font.size = Pt(10)
    work3_position_run.font.color.rgb = RGBColor(127, 140, 141)
    
    work3_content = right_cell.add_paragraph()
    work3_content.space_after = Pt(20)
    work3_content.paragraph_format.left_indent = Inches(0.2)
    work3_content_run = work3_content.add_run('• 见习学习：撰写教学设计20余篇；跟班见习，累计听课30余节；观摩优质赛课10余节')
    work3_content_run.font.size = Pt(9)
    work3_content_run.font.color.rgb = RGBColor(52, 73, 94)
    
    # 科研经历
    research_title = right_cell.add_paragraph()
    research_title.space_after = Pt(5)
    research_title_run = research_title.add_run('科研经历')
    research_title_run.font.size = Pt(16)
    research_title_run.font.bold = True
    research_title_run.font.color.rgb = RGBColor(44, 62, 80)
    
    # 科研经历下划线
    research_line = right_cell.add_paragraph()
    research_line.space_after = Pt(15)
    research_line_run = research_line.add_run('━' * 45)
    research_line_run.font.color.rgb = RGBColor(52, 152, 219)
    research_line_run.font.size = Pt(8)
    
    # 论文发表
    paper_title = right_cell.add_paragraph()
    paper_title.space_after = Pt(5)
    paper_title_run = paper_title.add_run('论文发表（1篇）')
    paper_title_run.font.size = Pt(12)
    paper_title_run.font.bold = True
    paper_title_run.font.color.rgb = RGBColor(44, 62, 80)
    
    paper_content1 = right_cell.add_paragraph()
    paper_content1.space_after = Pt(3)
    paper_content1.paragraph_format.left_indent = Inches(0.2)
    paper_content1_run = paper_content1.add_run('• 以独立作者在《西部教育论坛》发表"拔尖创新人才早期培养政策的回顾和展望"')
    paper_content1_run.font.size = Pt(9)
    paper_content1_run.font.color.rgb = RGBColor(52, 73, 94)
    
    paper_content2 = right_cell.add_paragraph()
    paper_content2.space_after = Pt(15)
    paper_content2.paragraph_format.left_indent = Inches(0.2)
    paper_content2_run = paper_content2.add_run('• 撰文《强国战略视域下拔尖创新人才政策的行动逻辑》（已投稿核心期刊）')
    paper_content2_run.font.size = Pt(9)
    paper_content2_run.font.color.rgb = RGBColor(52, 73, 94)
    
    # 课题研究
    project_title = right_cell.add_paragraph()
    project_title.space_after = Pt(5)
    project_title_run = project_title.add_run('课题研究（4项）')
    project_title_run.font.size = Pt(12)
    project_title_run.font.bold = True
    project_title_run.font.color.rgb = RGBColor(44, 62, 80)
    
    projects = [
        '参研四川省省级教育课题《基于"波普尔循环模型"的教师教学反思及创新能力培养循证教学改革》',
        '参研四川省省级教育课题《教师跨学科素养培养与教学实践创新研究》',
        '参研成都市市级教育课题《指向实践创新能力早期培养的小学跨学科项目式学习实践研究》'
    ]
    
    for project in projects:
        project_para = right_cell.add_paragraph()
        project_para.space_after = Pt(3)
        project_para.paragraph_format.left_indent = Inches(0.2)
        project_run = project_para.add_run(f'• {project}')
        project_run.font.size = Pt(9)
        project_run.font.color.rgb = RGBColor(52, 73, 94)
    
    # 学术活动
    activity_title = right_cell.add_paragraph()
    activity_title.space_before = Pt(15)
    activity_title.space_after = Pt(5)
    activity_title_run = activity_title.add_run('学术活动与获奖')
    activity_title_run.font.size = Pt(12)
    activity_title_run.font.bold = True
    activity_title_run.font.color.rgb = RGBColor(44, 62, 80)
    
    activities = [
        '"挑战杯"校级三等奖3次：跨学科教学素养调查研究、AI智能测评指导平台等',
        '"挑战杯"院级三等奖1次：托育养老综合服务体系发展策略研究',
        '"郑创汇"国际创新创业大赛院级决赛一等奖'
    ]
    
    for activity in activities:
        activity_para = right_cell.add_paragraph()
        activity_para.space_after = Pt(3)
        activity_para.paragraph_format.left_indent = Inches(0.2)
        activity_run = activity_para.add_run(f'• {activity}')
        activity_run.font.size = Pt(9)
        activity_run.font.color.rgb = RGBColor(52, 73, 94)
    
    # 底部个人宣言
    quote_para = right_cell.add_paragraph()
    quote_para.space_before = Pt(20)
    quote_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    quote_para.paragraph_format.left_indent = Inches(0.3)
    quote_para.paragraph_format.right_indent = Inches(0.3)
    quote_run = quote_para.add_run('坚信教育的本质是点亮与陪伴，我将用耐心与智慧激发每个孩子的潜能，与他们共同探索知识的乐趣。')
    quote_run.font.size = Pt(10)
    quote_run.font.italic = True
    quote_run.font.color.rgb = RGBColor(52, 152, 219)
    
    # 保存文档
    doc.save('何晓珊_专业简历.docx')
    print("✅ 专业简历已生成完成！")
    print("📄 文件名：何晓珊_专业简历.docx")

if __name__ == "__main__":
    try:
        create_professional_resume()
        print("\n🎨 设计特色：")
        print("   ▶ 采用经典的左右双栏布局")
        print("   ▶ 左栏深色背景配白色文字，专业大气")
        print("   ▶ 右栏白色背景配深色文字，清晰易读")
        print("   ▶ 使用专业的配色方案和层次化排版")
        print("   ▶ 头像区域预留，联系方式突出显示")
        print("   ▶ 内容结构化展示，重点信息一目了然")
    except ImportError:
        print("❌ 请先安装python-docx库：")
        print("   pip install python-docx")
    except Exception as e:
        print(f"❌ 生成简历时出错：{e}")